"""Build POC brief Word document for the AI Analyst system."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

OUT = Path("outputs/2026-04-29_AI_Analyst_POC_Brief.docx")
CHARTS = Path("outputs")

NAVY  = RGBColor(0x1A, 0x2E, 0x4A)
TEAL  = RGBColor(0x00, 0x87, 0x8A)
GOLD  = RGBColor(0xF0, 0xA8, 0x00)
LGREY = RGBColor(0xF4, 0xF5, 0xF7)
DGREY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _hex(rgb: RGBColor) -> str:
    return str(rgb)  # RGBColor.__str__ returns e.g. "1A2E4A"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tcPr.append(shd)


def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    pPr.append(shd)


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def add_run(para, text, bold=False, italic=False, colour=None, size=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if colour:
        run.font.color.rgb = colour
    if size:
        run.font.size = Pt(size)
    return run


def heading(doc, text, level=1, colour=NAVY, size=None):
    sizes = {1: 20, 2: 15, 3: 12}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    run = add_run(para, text, bold=True, colour=colour, size=size or sizes.get(level, 12))
    if level == 1:
        # Left rule via border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), _hex(TEAL))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def body(doc, text, space_after=4, italic=False, colour=DGREY, size=10.5):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, text, italic=italic, colour=colour, size=size)
    return para


def bullet(doc, text, bold_prefix=None, indent=0.3):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Inches(indent)
    para.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        add_run(para, bold_prefix + " ", bold=True, colour=NAVY, size=10.5)
        add_run(para, text, colour=DGREY, size=10.5)
    else:
        add_run(para, text, colour=DGREY, size=10.5)
    return para


def callout_box(doc, label, text, bg=LGREY, label_colour=TEAL):
    """Single-cell table styled as a callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)
    add_run(p1, label, bold=True, colour=label_colour, size=10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    add_run(p2, text, colour=DGREY, size=10)
    doc.add_paragraph()  # spacer


def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, h, bold=True, colour=WHITE, size=10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            if r_idx % 2 == 0:
                set_cell_bg(cell, RGBColor(0xF9, 0xFA, 0xFB))
            p = cell.paragraphs[0]
            add_run(p, val, colour=DGREY, size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)

    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


# ── Document sections ─────────────────────────────────────────────────────────

def cover_page(doc):
    # Title block
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(60)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, "AI Analyst — Proof of Concept Brief", bold=True, colour=NAVY, size=28)

    para2 = doc.add_paragraph()
    para2.paragraph_format.space_after = Pt(4)
    add_run(para2, "Demonstrating end-to-end AI-driven analytics: from raw data to executive report", colour=DGREY, size=13, italic=True)

    para3 = doc.add_paragraph()
    para3.paragraph_format.space_before = Pt(12)
    para3.paragraph_format.space_after = Pt(4)
    add_run(para3, "Prepared by:  ", bold=True, colour=DGREY, size=11)
    add_run(para3, "Analytics & Insights Team", colour=DGREY, size=11)

    para4 = doc.add_paragraph()
    para4.paragraph_format.space_after = Pt(4)
    add_run(para4, "Date:  ", bold=True, colour=DGREY, size=11)
    add_run(para4, "29 April 2026", colour=DGREY, size=11)

    para5 = doc.add_paragraph()
    para5.paragraph_format.space_after = Pt(4)
    add_run(para5, "Classification:  ", bold=True, colour=DGREY, size=11)
    add_run(para5, "Internal — Proof of Concept", colour=DGREY, size=11)

    doc.add_paragraph()

    callout_box(doc,
        "TL;DR",
        "In a single session, an AI analyst system took a raw member dataset, understood a business question, "
        "ran statistical analysis, produced four publication-quality charts, and assembled both a written report "
        "and an executive PowerPoint — without a human writing a single line of Python or spending hours in Excel. "
        "This document explains how, and what it means for the team.",
        bg=RGBColor(0xE8, 0xF4, 0xF5),
        label_colour=TEAL
    )

    page_break(doc)


def section_problem(doc):
    heading(doc, "1. The Problem This Solves")

    body(doc,
        "Today, analytics work at most organisations follows a predictable and expensive pattern: "
        "a business stakeholder asks a question, a data analyst manually extracts data, builds "
        "code or pivot tables, creates charts, writes up findings, formats slides — often over "
        "one to three days for a single question. The analyst's time is the bottleneck."
    )
    body(doc,
        "For a team covering superannuation member analytics at scale, this creates a backlog. "
        "Urgent questions get answered late. Ad-hoc exploratory questions often don't get answered "
        "at all. And when questions recur with new data, the process largely restarts from scratch."
    )

    heading(doc, "What if a question could be answered in minutes, not days?", level=2)

    body(doc,
        "The AI Analyst is a system built to answer that question. It is not a dashboard — "
        "dashboards answer the questions you anticipated. This system is designed to answer "
        "questions you haven't asked yet, on data you hand it today."
    )

    simple_table(doc,
        headers=["", "Traditional approach", "AI Analyst"],
        rows=[
            ["Time to first insight",      "1–3 days",             "Under 15 minutes"],
            ["Analyst hours per question", "4–12 hrs",             "~5 minutes of prompting"],
            ["Charts produced",            "1–3 (manual)",         "4+ (automated, publication quality)"],
            ["Written report",             "Separate task, hours", "Generated alongside analysis"],
            ["Slide deck",                 "Further hours",        "Auto-built from findings"],
            ["Repeatability",              "Manual re-run",        "Re-run with new data in seconds"],
            ["Skill required",             "Python / SQL / Excel", "Plain English question"],
        ],
        col_widths=[2.0, 2.2, 2.2]
    )

    page_break(doc)


def section_how(doc):
    heading(doc, "2. How the System Works")

    body(doc,
        "The AI Analyst is built on Claude Code — Anthropic's AI system for software and analytical "
        "work — configured with a set of specialist subagents and reusable skills. Think of it as a "
        "small analytics team that happens to operate entirely as software."
    )

    heading(doc, "The agent architecture", level=2)

    body(doc,
        "Five specialist agents collaborate, each with a distinct role. The orchestrator receives "
        "the business question and coordinates the rest:"
    )

    agent_rows = [
        ["Orchestrator",     "Receives the business question. Plans the work, delegates to specialists, inspects results, assembles the final deliverable. Never invents numbers."],
        ["Data Fetcher",     "Locates and loads the dataset. Runs a data quality profile. Flags nulls, duplicates, type issues. Saves a clean file for downstream use."],
        ["Analyst",          "Runs statistical analysis — segmentation, correlation, regression, hypothesis testing. Writes Python scripts. Outputs structured findings JSON."],
        ["Visualiser",       "Reads findings JSON and produces charts: bar charts, box plots, grouped comparisons. Saves PNGs with proper titles, axis labels, and source notes."],
        ["Narrative Writer", "Reads findings and charts. Produces the written report in BLUF structure (bottom line up front), calibrated for an executive non-technical audience."],
    ]
    simple_table(doc,
        headers=["Agent", "What it does"],
        rows=agent_rows,
        col_widths=[1.6, 4.8]
    )

    heading(doc, "Skills (reusable building blocks)", level=2)

    body(doc,
        "Agents draw on a library of reusable skills — pre-defined capabilities invoked on demand:"
    )
    bullet(doc, "data profiling — shape, schema, null rates, duplicate detection, cardinality flags", bold_prefix="Data Profiling:")
    bullet(doc, "SQL query building — dialect-aware, tested query construction for warehouse extraction", bold_prefix="SQL Builder:")
    bullet(doc, "chart building — chart-type selection, styling standards, matplotlib/plotly templates", bold_prefix="Chart Builder:")
    bullet(doc, "narrative writing — BLUF structure, executive tone, Australian English conventions", bold_prefix="Insight Narrative:")

    doc.add_paragraph()

    heading(doc, "The workflow in plain English", level=2)

    steps = [
        ("1. Profile", "Hand the system a dataset. It runs an automated health check: shape, nulls, duplicates, distributions, quality flags."),
        ("2. Question", "Ask a business question in plain English. The orchestrator restates it, confirms scope, and writes a plan."),
        ("3. Fetch & clean", "The data fetcher loads the data, removes duplicates, handles nulls, saves a clean file."),
        ("4. Analyse", "The analyst runs the statistical work — in this case, participation rates, conditional amounts, regression, cross-tabs by salary quartile."),
        ("5. Visualise", "The visualiser produces charts directly from the analysis output — no copy-paste, no manual formatting."),
        ("6. Write up", "The narrative writer produces a structured markdown report: exec summary, key findings, caveats."),
        ("7. Deliver", "The orchestrator assembles the final deliverable. Optionally: export to PowerPoint for sharing."),
    ]
    for label, desc in steps:
        bullet(doc, desc, bold_prefix=label)

    doc.add_paragraph()
    page_break(doc)


def section_demo(doc):
    heading(doc, "3. The Demo: What We Did Today")

    body(doc,
        "To demonstrate the system, we ran a real analytical question against a superannuation member "
        "dataset — the kind of question a product or member experience team might ask any week."
    )

    callout_box(doc,
        "The question",
        "\"What's driving differences in salary sacrifice behaviour across age bands?\"",
        bg=RGBColor(0xE8, 0xF4, 0xF5), label_colour=TEAL
    )

    heading(doc, "The dataset", level=2)
    simple_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["File",             "super_members_2026Q1.csv"],
            ["Snapshot date",    "31 March 2026"],
            ["Rows (raw)",       "5,025"],
            ["Rows (clean)",     "5,000 (25 exact duplicates removed)"],
            ["Columns",          "20"],
            ["Key variables",    "age_band, salary_sacrifice_12mo_aud, salary_aud, balance_aud, tenure_years, join_channel, investment_option, last_login_days_ago, has_app_installed"],
        ],
        col_widths=[1.8, 4.6]
    )

    heading(doc, "What the system did — step by step", level=2)

    body(doc, "Step 1 — Data profile (seconds)")
    body(doc,
        "The data-profiling skill ran automatically: detected 25 duplicate rows, flagged salary_aud "
        "as having 100 nulls (2%), noted snapshot_date as a constant column (single point-in-time), "
        "and identified that salary_sacrifice_12mo_aud had a median of $0 across all members — "
        "meaning most members make no voluntary contributions at all. This insight directly shaped "
        "how the analysis framed participation vs amount as two separate dimensions.",
        italic=False, size=10.5
    )

    body(doc, "Step 2 — Orchestration plan (seconds)")
    body(doc,
        "The orchestrator restated the question, identified the relevant dimensions "
        "(participation rate by age band; conditional amount; co-variates including salary, tenure, "
        "engagement, gender; interaction effects), and delegated sequentially to the specialist agents.",
        size=10.5
    )

    body(doc, "Step 3 — Analysis (minutes)")
    body(doc,
        "The analyst agent wrote and executed Python scripts covering: participation rate by age band "
        "with confidence intervals, conditional median and mean sacrifice amounts, cross-tabulation "
        "by salary quartile, linear probability regression (OLS) with standardised coefficients, "
        "sacrifice-as-percentage-of-salary by age band, and sub-group breakdowns by digital engagement "
        "and gender.",
        size=10.5
    )

    body(doc, "Step 4 — Charts (minutes)")
    body(doc,
        "The visualiser produced four publication-quality charts from the findings JSON — no manual "
        "formatting, no copy-paste from Python output:",
        size=10.5
    )
    bullet(doc, "Participation rate by age band (bar chart, with overall average reference line)")
    bullet(doc, "Conditional sacrifice amount — median and mean (grouped bar by age band)")
    bullet(doc, "Participation by age band × salary quartile (grouped bar, cross-tab heat)")
    bullet(doc, "Sacrifice as % of salary by age band (box plot, participants only)")

    body(doc, "Step 5 — Written report (minutes)")
    body(doc,
        "The narrative writer produced a 12-page structured report: executive summary (5 bullets), "
        "six substantive sections with embedded chart references, an implications section with "
        "four cohort-specific recommendations, and a caveats section. All in Australian English, "
        "all citing specific numbers from the analysis.",
        size=10.5
    )

    body(doc, "Step 6 — PowerPoint deck (minutes)")
    body(doc,
        "A Python script assembled an 8-slide branded executive deck: title slide, executive summary, "
        "one slide per key finding with the chart embedded and an annotation panel, an opportunities "
        "slide with four cohort cards, and a caveats slide. The deck is self-contained and ready to send.",
        size=10.5
    )

    doc.add_paragraph()
    callout_box(doc,
        "Total elapsed time",
        "From raw CSV to a complete executive PowerPoint with four charts and a written report: "
        "approximately 30–40 minutes of wall clock time, with under 5 minutes of human input "
        "(two prompts). No Python written by hand. No Excel. No manual chart formatting.",
        bg=RGBColor(0xE8, 0xF4, 0xF5), label_colour=TEAL
    )

    page_break(doc)


def section_findings(doc):
    heading(doc, "4. What the Analysis Found")

    body(doc,
        "To give a concrete sense of output quality, here are the headline findings the AI Analyst produced. "
        "These came directly from the automated analysis — no human validated or modified them before inclusion "
        "in the report."
    )

    findings = [
        ("The age gap is a participation problem, not an amount problem.",
         "Participation climbs from 9.2% (18–24 year-olds) to 63.5% (55–64 year-olds). "
         "But once a member opts in, the annual amount is nearly flat across all age bands "
         "— median $6,024 to $7,396. Strategy should focus on getting members to start, not on "
         "getting existing participants to contribute more."),
        ("Age predicts participation more strongly than salary.",
         "A linear regression showed age rank has a standardised coefficient of 0.110 — more than "
         "twice that of log salary (0.048). Even after controlling for what members earn, older members "
         "participate at substantially higher rates. Life-stage factors (proximity to retirement, "
         "financial advice, discharged family obligations) are driving the effect."),
        ("For younger members, salary is the gating factor.",
         "In the 18–24 band, the highest salary quartile reaches only 33% participation — still "
         "low by any measure, but compared to near-zero in Q1. Salary growth matters most for "
         "activating this cohort. Targeted communications at salary milestones are likely to outperform "
         "generic campaigns."),
        ("An unexpected reversal in older high-earners.",
         "In the 55–64 and 65+ bands, the highest salary quartile actually participates less than Q2/Q3. "
         "Non-participants in those bands have higher median salaries than participants. This likely "
         "reflects the concessional contributions cap ($30K p.a. including employer SG) becoming "
         "binding, or alternative contribution strategies being used by high-income older members."),
        ("Digital engagement does not predict salary sacrifice.",
         "Having the app installed and login recency are effectively uncorrelated with whether a "
         "member participates. In several age bands, app-installed members participate at lower rates. "
         "The decision to salary sacrifice appears driven by financial capacity and life-stage "
         "awareness — not platform engagement."),
        ("Biggest single opportunity: 35–44 non-participants.",
         "718 members in the 35–44 band are earning meaningful salaries and not participating, "
         "despite 48% of their peers doing so. This cohort sits at the inflection point of the "
         "participation curve and is most likely to respond to peer comparison and default opt-in nudges."),
    ]

    for i, (title, desc) in enumerate(findings):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)
        add_run(para, f"{i+1}.  {title}", bold=True, colour=NAVY, size=11)
        body(doc, desc, space_after=4, size=10.5)

    page_break(doc)


def section_deliverables(doc):
    heading(doc, "5. Deliverables Produced in This Session")

    body(doc,
        "Every file below was generated automatically by the AI Analyst system. No human "
        "wrote code, formatted a chart, or drafted prose for any of these outputs."
    )

    simple_table(doc,
        headers=["File", "Type", "Description"],
        rows=[
            ["_profile_super_members_2026Q1.md",        "Markdown",   "Automated data quality profile: shape, schema, null rates, flags"],
            ["2026-04-29_findings.json",                 "JSON",       "Machine-readable results: participation rates, amounts, regression coefficients, cross-tabs"],
            ["2026-04-29_participation_rate_by_age_band.png",      "Chart (PNG)", "Bar chart — participation rate by age band"],
            ["2026-04-29_conditional_amount_by_age_band.png",      "Chart (PNG)", "Grouped bar — median/mean sacrifice amount (participants only)"],
            ["2026-04-29_participation_by_age_salary_quartile.png","Chart (PNG)", "Grouped bar — participation by age band × salary quartile"],
            ["2026-04-29_sacrifice_pct_salary_by_age.png",         "Chart (PNG)", "Box plot — sacrifice as % of salary by age band"],
            ["2026-04-29_salary_sacrifice_analysis.md",  "Markdown",   "Full written report: exec summary, 6 finding sections, implications, caveats"],
            ["2026-04-29_salary_sacrifice_report.pptx",  "PowerPoint", "8-slide executive deck with embedded charts, annotation panels, opportunity cards"],
            ["2026-04-29_AI_Analyst_POC_Brief.docx",     "Word",       "This document"],
        ],
        col_widths=[2.8, 1.2, 2.4]
    )

    page_break(doc)


def section_next(doc):
    heading(doc, "6. What This Could Look Like at Scale")

    body(doc,
        "This demonstration used one question and one dataset. The system is designed to generalise. "
        "Below are use cases where this approach could add direct value for the Analytics & Insights "
        "function and for the broader business."
    )

    heading(doc, "Near-term use cases", level=2)

    use_cases = [
        ("Member segmentation on demand",
         "Any time a product manager or executive asks \"tell me about members who X\" — "
         "churn risk, low engagement, approaching retirement, zero beneficiary nomination — "
         "the system can produce a segmentation report without analyst queuing time."),
        ("Investment option migration analysis",
         "Understanding why members switch options, or which cohorts are most likely to switch, "
         "is a recurring question. The system can re-run this analysis each quarter with fresh data."),
        ("Fee and balance benchmarking",
         "Cross-cohort comparisons of fee drag, balance trajectory, and contribution adequacy "
         "can be automated against any new data snapshot."),
        ("Campaign post-analysis",
         "After a member communications campaign, the system can run a before/after "
         "comparison on the target cohort vs a control group automatically."),
        ("Contribution adequacy reporting",
         "Checking whether members across demographic cohorts are on track for an adequate "
         "retirement balance — a recurring compliance and member outcomes question."),
    ]
    for title, desc in use_cases:
        bullet(doc, desc, bold_prefix=title)

    heading(doc, "Longer-term possibilities", level=2)

    bullet(doc, "Scheduled weekly or monthly analysis runs triggered automatically when new data lands in the warehouse")
    bullet(doc, "Integration with Snowflake or BigQuery — the data fetcher already supports SQL query building against warehouse dialects")
    bullet(doc, "A self-serve interface where non-technical team members can ask questions in plain English and receive a report")
    bullet(doc, "Audit trail by design — every analysis produces a findings JSON and a Python script, making results reproducible and reviewable")

    heading(doc, "What it is not (important caveats)", level=2)
    body(doc,
        "This is a decision-support tool, not a replacement for analytical judgement. The system "
        "does not know when a finding is commercially sensitive, politically complex, or requires "
        "context only a human holds. Every output should be reviewed before being shared externally "
        "or used in a regulatory context. The system is a force multiplier for the analyst — it "
        "eliminates the mechanical work, not the thinking."
    )

    page_break(doc)


def section_tech(doc):
    heading(doc, "7. Technical Stack")

    body(doc, "For those who want to understand what is running under the hood:")

    simple_table(doc,
        headers=["Component", "Technology", "Notes"],
        rows=[
            ["AI model",          "Claude Sonnet 4.6 (Anthropic)", "Powers all agents; runs locally via Claude Code CLI"],
            ["Orchestration",     "Claude Code multi-agent framework", "Agents defined as markdown files in .claude/agents/"],
            ["Data analysis",     "Python 3.11 — pandas, numpy, scipy", "Scripts written and executed by the analyst agent"],
            ["Visualisation",     "matplotlib, seaborn",               "Charts written and rendered by the visualiser agent"],
            ["Report generation", "Markdown → python-docx / python-pptx", "Converted to Word/PowerPoint by a build script"],
            ["Data storage",      "Local filesystem + parquet",        "Clean datasets saved to data/; outputs to outputs/"],
            ["Skills",            "Reusable markdown skill packages",  "data-profiling, chart-builder, insight-narrative, sql-query-builder"],
            ["Version control",   "Git",                               "All agent definitions, scripts, and skills under source control"],
        ],
        col_widths=[1.6, 2.2, 2.6]
    )

    heading(doc, "Where the code lives", level=2)
    body(doc, "Everything is in the ai-analyst project directory:")
    bullet(doc, "Agent definitions in .claude/agents/ — each a markdown file describing the agent's role, tools, and behaviour")
    bullet(doc, "Skill packages in .claude/skills/ — reusable instruction sets invoked by agents")
    bullet(doc, "Analysis scripts written by the agents in scripts/ — reproducible, version-controlled")
    bullet(doc, "All outputs in outputs/ — datestamped filenames for auditability")

    doc.add_paragraph()
    body(doc,
        "Standing up the system for a new dataset or a new question requires no code changes — "
        "only a new data file and a new prompt.",
        italic=True, colour=DGREY
    )

    page_break(doc)


def section_conclusion(doc):
    heading(doc, "8. Summary and Recommended Next Steps")

    callout_box(doc,
        "What this proves",
        "A multi-agent AI system can take a raw dataset and a plain-English business question "
        "and produce executive-grade analytical output — including statistical analysis, "
        "publication-quality charts, a written report, and a slide deck — in under an hour, "
        "with minimal human effort. The output quality is comparable to what a skilled analyst "
        "would produce manually in one to three days.",
        bg=RGBColor(0xE8, 0xF4, 0xF5), label_colour=TEAL
    )

    heading(doc, "Recommended next steps", level=2)

    steps = [
        ("Review the outputs",
         "Read the written report and flip through the PowerPoint. Assess whether the analytical "
         "depth and output quality would meet the bar for internal reporting."),
        ("Identify two or three test questions",
         "Pick real questions from the current backlog — ones that have been delayed due to "
         "analyst capacity — and run them through the system. Compare the outputs to what "
         "would have been produced manually."),
        ("Evaluate data access",
         "The current system reads files from a local directory. The logical next step is "
         "connecting it to the data warehouse (Snowflake / BigQuery) so it can pull data directly "
         "without a manual extract step."),
        ("Define the governance model",
         "Agree on a review process: who checks AI-produced analysis before it goes to stakeholders, "
         "what gets logged, and how findings are attributed. This is lightweight to set up but "
         "important to have before broader rollout."),
        ("Consider scheduling",
         "High-value recurring analyses (quarterly member cohort reports, contribution adequacy "
         "snapshots) could be scheduled to run automatically when new data arrives."),
    ]
    for title, desc in steps:
        bullet(doc, desc, bold_prefix=title)

    doc.add_paragraph()
    body(doc,
        "Questions or feedback: contact your Analytics & Insights team. "
        "The system is running in the ai-analyst project and can be demonstrated live.",
        italic=True, colour=DGREY, size=10
    )


# ── Assemble ──────────────────────────────────────────────────────────────────

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DGREY
    style.paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    set_margins(doc)
    set_default_style(doc)

    cover_page(doc)
    section_problem(doc)
    section_how(doc)
    section_demo(doc)
    section_findings(doc)
    section_deliverables(doc)
    section_next(doc)
    section_tech(doc)
    section_conclusion(doc)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
