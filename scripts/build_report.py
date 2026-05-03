"""
Parametric Word report builder — consultant-quality output for any analysis.

Usage:
    python scripts/build_report.py outputs/report_config.json
    python scripts/build_report.py  # builds salary sacrifice demo

Design philosophy:
- Clean, minimal — generous whitespace, precise typography
- Brand-consistent colour system
- Professional table and callout styles
- Proper running headers/footers
- No inline formatting inconsistencies
"""

from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import json
import sys
import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── Colour system ─────────────────────────────────────────────────────────────

class C:
    NAVY    = RGBColor(0x1B, 0x35, 0x57)
    TEAL    = RGBColor(0x0A, 0x8A, 0x8F)
    AMBER   = RGBColor(0xE8, 0xA0, 0x20)
    CHARCOAL = RGBColor(0x2C, 0x3E, 0x50)
    SLATE   = RGBColor(0x8A, 0x99, 0xAA)
    LGREY   = RGBColor(0xF7, 0xF8, 0xFA)
    MGREY   = RGBColor(0xE5, 0xE8, 0xEC)
    WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
    TEAL_BG = RGBColor(0xE5, 0xF5, 0xF6)
    AMBER_BG = RGBColor(0xFD, 0xF3, 0xE0)
    POSITIVE = RGBColor(0x18, 0x7A, 0x4E)
    NEGATIVE = RGBColor(0xC0, 0x39, 0x2B)

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"

# ── XML helpers ───────────────────────────────────────────────────────────────

def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(rgb))
    tcPr.append(shd)


def _set_para_border_left(para, rgb: RGBColor, sz: int = 24):
    """Add a coloured left border (rule) to a paragraph — callout style."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), _hex(rgb))
    pBdr.append(left)
    pPr.append(pBdr)


def _set_bottom_border(para, rgb: RGBColor, sz: int = 6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _hex(rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_col_width(table, col_idx: int, width_inches: float):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── Document primitives ───────────────────────────────────────────────────────

def _run(para, text: str, *, bold=False, italic=False,
         colour: RGBColor = None, size: float = None, font: str = FONT_BODY):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if colour:
        run.font.color.rgb = colour
    if size:
        run.font.size = Pt(size)
    return run


def _spacer(doc: Document, space_before: float = 0, space_after: float = 6) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)


# ── Heading system ────────────────────────────────────────────────────────────

def h1(doc: Document, text: str, numbered: str = None) -> None:
    """Section heading with teal bottom rule."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(4)
    if numbered:
        _run(para, f"{numbered}  ", bold=True, colour=C.TEAL, size=15)
    _run(para, text, bold=True, colour=C.NAVY, size=16, font=FONT_HEAD)
    _set_bottom_border(para, C.TEAL, sz=8)


def h2(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(3)
    _run(para, text, bold=True, colour=C.NAVY, size=12.5, font=FONT_HEAD)


def h3(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    _run(para, text, bold=True, colour=C.TEAL, size=11, font=FONT_HEAD)


def body(doc: Document, text: str, *,
         space_after: float = 5, italic: bool = False,
         colour: RGBColor = None, size: float = 10.5) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(15)
    _run(para, text, italic=italic, colour=colour or C.CHARCOAL, size=size)


def bullet(doc: Document, text: str, *,
           bold_prefix: str = None, level: int = 0) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = Pt(14)
    if bold_prefix:
        _run(para, bold_prefix + "  ", bold=True, colour=C.NAVY, size=10.5)
        _run(para, text, colour=C.CHARCOAL, size=10.5)
    else:
        _run(para, text, colour=C.CHARCOAL, size=10.5)


# ── Callout styles ────────────────────────────────────────────────────────────

def callout(doc: Document, label: str, text: str,
            style: str = "teal") -> None:
    """
    Styled callout box with coloured left border rule.
    style: "teal" | "amber" | "navy" | "insight"
    """
    colour_map = {
        "teal": (C.TEAL, C.TEAL_BG),
        "amber": (C.AMBER, C.AMBER_BG),
        "navy": (C.NAVY, C.LGREY),
        "insight": (C.POSITIVE, RGBColor(0xEA, 0xF7, 0xF0)),
    }
    border_colour, bg_colour = colour_map.get(style, (C.TEAL, C.TEAL_BG))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    _set_cell_bg(cell, bg_colour)

    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_before = Pt(7)
    p_label.paragraph_format.space_after = Pt(3)
    p_label.paragraph_format.left_indent = Inches(0.15)
    _set_para_border_left(p_label, border_colour, sz=28)
    _run(p_label, label.upper(), bold=True, colour=border_colour, size=9.5)

    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(8)
    p_body.paragraph_format.left_indent = Inches(0.15)
    _set_para_border_left(p_body, border_colour, sz=28)
    _run(p_body, text, colour=C.CHARCOAL, size=10.5)

    _spacer(doc, space_after=8)


def pull_quote(doc: Document, text: str, attribution: str = "") -> None:
    """Large pull quote — key finding as visual break."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.left_indent = Inches(0.4)
    para.paragraph_format.right_indent = Inches(0.4)
    _set_para_border_left(para, C.TEAL, sz=40)
    _run(para, text, italic=True, colour=C.NAVY, size=14)
    if attribution:
        para2 = doc.add_paragraph()
        para2.paragraph_format.left_indent = Inches(0.4)
        para2.paragraph_format.space_after = Pt(10)
        _run(para2, f"— {attribution}", colour=C.SLATE, size=10, italic=True)


# ── Table builder ─────────────────────────────────────────────────────────────

def data_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] = None,
               caption: str = None) -> None:
    """Premium table — navy header, alternating rows, no heavy borders."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for c, hdr in enumerate(headers):
        cell = table.cell(0, c)
        _set_cell_bg(cell, C.NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        _run(p, hdr, bold=True, colour=C.WHITE, size=9.5)

    # Data rows
    for r, row_data in enumerate(rows):
        bg = C.LGREY if r % 2 == 0 else C.WHITE
        for c, val in enumerate(row_data):
            cell = table.cell(r + 1, c)
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _run(p, str(val), colour=C.CHARCOAL, size=9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            _set_col_width(table, i, w)

    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(8)
        _run(cap, caption, italic=True, colour=C.SLATE, size=9)
    else:
        _spacer(doc, space_after=8)


def kpi_row(doc: Document, metrics: list[dict]) -> None:
    """
    Horizontal KPI strip — up to 4 metrics side by side.
    metric dict: {"label": str, "value": str, "note": str (opt), "flag": "positive"|"negative"|None}
    """
    n = min(len(metrics), 4)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, m in enumerate(metrics[:n]):
        cell = table.cell(0, i)
        _set_cell_bg(cell, C.LGREY)
        # Value
        p_val = cell.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(6)
        p_val.paragraph_format.space_after = Pt(1)
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        flag_colour = C.POSITIVE if m.get("flag") == "positive" else (
            C.NEGATIVE if m.get("flag") == "negative" else C.NAVY)
        _run(p_val, m["value"], bold=True, colour=flag_colour, size=22)

        # Label
        p_lbl = cell.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_after = Pt(2)
        _run(p_lbl, m["label"], colour=C.CHARCOAL, size=9)

        # Note
        if m.get("note"):
            p_note = cell.add_paragraph()
            p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_note.paragraph_format.space_after = Pt(6)
            _run(p_note, m["note"], italic=True, colour=C.SLATE, size=8.5)

    _spacer(doc, space_after=10)


# ── Document structure ────────────────────────────────────────────────────────

def cover_page(doc: Document, cfg: ReportConfig) -> None:
    """Premium cover with navy header band and structured metadata."""

    # Top colour band (simulated via borderless table)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, C.NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    _run(p, "  " + cfg.classification,
         colour=RGBColor(0xB0, 0xCC, 0xE0), size=10, italic=True)

    _spacer(doc, space_before=28)

    # Title block
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(6)
    _run(title_para, cfg.title, bold=True, colour=C.NAVY, size=30, font=FONT_HEAD)

    if cfg.subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(16)
        _run(sub, cfg.subtitle, italic=True, colour=C.SLATE, size=14)

    # Teal rule
    rule = doc.add_paragraph()
    _set_bottom_border(rule, C.TEAL, sz=24)
    rule.paragraph_format.space_after = Pt(14)

    # Metadata block
    meta_items = [
        ("Prepared by", cfg.prepared_by),
        ("Date", cfg.prepared_date or datetime.date.today().strftime("%d %B %Y")),
        ("Organisation", cfg.organisation) if cfg.organisation else None,
        ("Data source", cfg.data_source) if cfg.data_source else None,
        ("Period covered", cfg.period_covered) if cfg.period_covered else None,
        ("Classification", cfg.classification),
    ]
    for item in meta_items:
        if not item:
            continue
        label, value = item
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        _run(para, f"{label}:  ", bold=True, colour=C.NAVY, size=10.5)
        _run(para, value, colour=C.CHARCOAL, size=10.5)

    _spacer(doc, space_before=14)

    # Bottom line up front callout
    if cfg.bottom_line:
        callout(doc, "Bottom line", cfg.bottom_line, style="teal")

    doc.add_page_break()


def exec_summary_page(doc: Document, bullets: list[str]) -> None:
    h1(doc, "Executive Summary")
    for b in bullets:
        bullet(doc, b)
    _spacer(doc, space_after=6)
    doc.add_page_break()


def findings_section(doc: Document, section_num: str, section_title: str,
                     findings: list[dict]) -> None:
    """
    findings: list of dicts with keys:
        title, body, chart_path (opt), chart_caption (opt),
        table (opt: {headers, rows, col_widths}),
        kpis (opt: list of metric dicts),
        callout (opt: {label, text, style}),
        bullets (opt: list of str)
    """
    h1(doc, section_title, numbered=section_num)

    for finding in findings:
        if finding.get("title"):
            h2(doc, finding["title"])

        if finding.get("kpis"):
            kpi_row(doc, finding["kpis"])

        if finding.get("body"):
            body(doc, finding["body"])

        if finding.get("bullets"):
            for b in finding["bullets"]:
                bullet(doc, b)

        if finding.get("chart_path"):
            p = Path(finding["chart_path"])
            if p.exists():
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(4)
                    run = para.add_run()
                    run.add_picture(str(p), width=Inches(6.0))
                except Exception:
                    pass
            if finding.get("chart_caption"):
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(8)
                _run(cap, finding["chart_caption"],
                     italic=True, colour=C.SLATE, size=9)

        if finding.get("table"):
            t = finding["table"]
            data_table(doc, t["headers"], t["rows"],
                       t.get("col_widths"), t.get("caption"))

        if finding.get("callout"):
            c = finding["callout"]
            callout(doc, c["label"], c["text"], style=c.get("style", "teal"))

        _spacer(doc, space_after=6)


def recommendations_section(doc: Document, section_num: str,
                             recommendations: list[dict]) -> None:
    """
    recommendations: list of dicts:
        {title, body, owner (opt), timeframe (opt), priority (opt: "high"|"medium"|"low")}
    """
    h1(doc, "Recommendations", numbered=section_num)

    if recommendations:
        headers = ["#", "Recommendation", "Owner", "Timeframe", "Priority"]
        rows = []
        for i, r in enumerate(recommendations, 1):
            rows.append([
                str(i),
                r["title"],
                r.get("owner", "—"),
                r.get("timeframe", "—"),
                r.get("priority", "—").capitalize(),
            ])
        data_table(doc, headers, rows, col_widths=[0.3, 3.0, 1.2, 1.0, 0.8])

    for i, rec in enumerate(recommendations, 1):
        h3(doc, f"{i}. {rec['title']}")
        body(doc, rec["body"])
        _spacer(doc)


def caveats_section(doc: Document, section_num: str, caveats: list[str]) -> None:
    h1(doc, "Caveats & Limitations", numbered=section_num)
    for c in caveats:
        bullet(doc, c)
    _spacer(doc)


def appendix_section(doc: Document, items: list[dict]) -> None:
    """
    items: list of dicts: {title, content (str or list), type: "text"|"bullets"|"table"}
    """
    doc.add_page_break()
    h1(doc, "Appendix")
    for item in items:
        h2(doc, item["title"])
        if item.get("type") == "bullets" or isinstance(item.get("content"), list):
            for b in item["content"]:
                bullet(doc, b)
        elif item.get("type") == "table":
            t = item["content"]
            data_table(doc, t["headers"], t["rows"], t.get("col_widths"))
        else:
            body(doc, str(item.get("content", "")))
        _spacer(doc)


def add_running_headers_footers(doc: Document, cfg: ReportConfig) -> None:
    """Add header (doc title) and footer (page number + classification)."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.paragraph_format.space_after = Pt(0)
        _set_bottom_border(hp, C.MGREY, sz=4)
        _run(hp, cfg.organisation + "  |  " if cfg.organisation else "",
             colour=C.SLATE, size=9)
        _run(hp, cfg.title, bold=True, colour=C.NAVY, size=9)
        tab_run = hp.add_run()
        tab_run.font.size = Pt(9)
        tab_run.font.color.rgb = C.SLATE
        tab_run.font.name = FONT_BODY

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.paragraph_format.space_before = Pt(0)
        _run(fp, cfg.classification + "  |  Page ", colour=C.SLATE, size=8.5, italic=True)
        # Page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_xml = fp.add_run()
        run_xml.font.size = Pt(8.5)
        run_xml.font.color.rgb = C.SLATE
        run_xml._r.append(fldChar1)
        run_xml._r.append(instrText)
        run_xml._r.append(fldChar2)
        _run(fp, f"  |  {cfg.prepared_date}",
             colour=C.SLATE, size=8.5, italic=True)


# ── Config and assembly ───────────────────────────────────────────────────────

@dataclass
class ReportConfig:
    title: str
    subtitle: str = ""
    organisation: str = ""
    prepared_by: str = "Analytics & Insights"
    prepared_date: str = ""
    classification: str = "CONFIDENTIAL — INTERNAL USE ONLY"
    data_source: str = ""
    period_covered: str = ""
    bottom_line: str = ""
    contact: str = ""

    exec_summary: list[str] = field(default_factory=list)
    sections: list[dict] = field(default_factory=list)
    recommendations: list[dict] = field(default_factory=list)
    caveats: list[str] = field(default_factory=list)
    appendix: list[dict] = field(default_factory=list)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.font.color.rgb = C.CHARCOAL
    style.paragraph_format.space_after = Pt(5)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.2)


def build_report(cfg: ReportConfig, out_path: Path) -> None:
    if not cfg.prepared_date:
        cfg.prepared_date = datetime.date.today().strftime("%d %B %Y")

    doc = Document()
    set_doc_defaults(doc)

    cover_page(doc, cfg)

    if cfg.exec_summary:
        exec_summary_page(doc, cfg.exec_summary)

    for i, section in enumerate(cfg.sections, 1):
        findings_section(
            doc,
            section_num=str(i),
            section_title=section["title"],
            findings=section.get("findings", []),
        )

    if cfg.recommendations:
        recommendations_section(doc, str(len(cfg.sections) + 1), cfg.recommendations)

    if cfg.caveats:
        caveats_section(doc, str(len(cfg.sections) + 2), cfg.caveats)

    if cfg.appendix:
        appendix_section(doc, cfg.appendix)

    add_running_headers_footers(doc, cfg)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Report saved: {out_path}")


def build_from_json(config_path: str) -> None:
    with open(config_path) as f:
        raw = json.load(f)
    cfg = ReportConfig(**{k: v for k, v in raw.items() if k != "output_path"})
    out = Path(raw.get("output_path", "outputs/report.docx"))
    build_report(cfg, out)


# ── Salary sacrifice demo report ─────────────────────────────────────────────

def build_salary_sacrifice_report() -> None:
    cfg = ReportConfig(
        title="Salary Sacrifice Behaviour by Age Band",
        subtitle="Superannuation Member Analysis — Q1 2026",
        organisation="Analytics & Insights",
        prepared_by="Bharath",
        prepared_date="29 April 2026",
        classification="CONFIDENTIAL — INTERNAL USE ONLY",
        data_source="super_members_2026Q1.csv",
        period_covered="31 March 2026 (point-in-time snapshot)",
        bottom_line=(
            "The 54-percentage-point participation gap between our youngest and oldest members is "
            "almost entirely a question of who opts in, not how much they contribute once they do. "
            "The 35–44 cohort represents the single largest near-term opportunity: 718 non-participants "
            "at the participation inflection point, with both the salary and the time horizon to benefit."
        ),

        exec_summary=[
            "Participation climbs from 9.2% (age 18–24) to 63.5% (age 55–64) — a 54-percentage-point gap driven by life stage, not salary.",
            "Once a member opts in, annual sacrifice amounts are nearly flat across all age bands ($6,024–$7,396 median).",
            "Age is 2.3× more predictive of participation than salary: standardised OLS coefficients 0.110 vs 0.048.",
            "High-earners aged 55–64 participate less than mid-earners — likely reflecting the $30K concessional cap becoming binding.",
            "718 non-participating 35–44 year-olds represent the single largest activation opportunity in the fund.",
        ],

        sections=[
            {
                "title": "The Core Finding: A Participation Gap, Not an Amount Gap",
                "findings": [
                    {
                        "title": "Participation rises sharply with age — amounts do not",
                        "kpis": [
                            {"value": "48.7%", "label": "Fund participation rate", "note": "2,434 of 5,000 members"},
                            {"value": "9.2%", "label": "18–24 participation", "note": "n=31 of 338", "flag": "negative"},
                            {"value": "63.5%", "label": "55–64 participation", "note": "n=475 of 748", "flag": "positive"},
                            {"value": "$6,300", "label": "Median annual sacrifice", "note": "Participants only, all ages"},
                        ],
                        "body": (
                            "The fund's participation rate ranges from 9.2% among 18–24 year-olds to 63.5% among the 55–64 cohort — "
                            "a 54-percentage-point gap that grows almost monotonically with age. Yet once a member begins salary sacrificing, "
                            "the annual amount they contribute is remarkably stable across all age cohorts: the conditional median ranges "
                            "narrowly from $6,024 (35–44 band) to $7,396 (18–24 band, n=31 — treat as indicative). "
                            "The strategic implication is clear: the activation gap, not the generosity gap, is where intervention will generate returns."
                        ),
                        "chart_path": "outputs/2026-04-29_participation_rate_by_age_band.png",
                        "chart_caption": "Figure 1: Participation rate by age band. Fund average 48.7% (dashed reference line). n=5,000 members.",
                        "table": {
                            "headers": ["Age Band", "Members", "Participants", "Rate", "Median Amount (AUD)", "n (participants)"],
                            "rows": [
                                ["18–24",   "338",    "31",    "9.2%",  "$7,396", "31"],
                                ["25–34",   "787",    "213",   "27.1%", "$6,301", "213"],
                                ["35–44",   "1,382",  "664",   "48.0%", "$6,024", "664"],
                                ["45–54",   "1,448",  "876",   "60.5%", "$6,086", "876"],
                                ["55–64",   "748",    "475",   "63.5%", "$6,263", "475"],
                                ["65+",     "297",    "175",   "58.9%", "$6,298", "175"],
                                ["Total",   "5,000",  "2,434", "48.7%", "$6,198", "2,434"],
                            ],
                            "col_widths": [0.9, 0.9, 1.1, 0.8, 1.6, 1.3],
                            "caption": "Table 1: Participation and conditional sacrifice amounts by age band. Source: super_members_2026Q1.csv.",
                        },
                    },
                ],
            },
            {
                "title": "What Drives Participation: Age Dominates Salary",
                "findings": [
                    {
                        "title": "Age is 2.3× more predictive than salary",
                        "body": (
                            "A linear probability model (OLS) regressing participation on age rank, log salary, and tenure finds "
                            "that age rank carries a standardised coefficient of 0.110 — more than twice the coefficient on log salary (0.048). "
                            "Salary alone explains 3.8% of participation variance (R²); adding age rank and tenure raises this to 9.5%. "
                            "The large unexplained variance (90.5%) suggests that life-stage factors not captured in this dataset — "
                            "proximity to retirement, family obligations discharged, financial advice relationships — are the dominant drivers."
                        ),
                        "chart_path": "outputs/2026-04-29_participation_by_age_salary_quartile.png",
                        "chart_caption": "Figure 2: Participation rate by age band and salary quartile. Within-band salary effect largest for younger cohorts.",
                        "callout": {
                            "label": "Key insight",
                            "text": "For the 55–64 band, top-quartile earners (Q4) participate at lower rates than Q2 and Q3. Non-participants in this band have higher median salaries ($129,200) than participants ($111,400). This reversal likely reflects the $30,000 concessional cap (including employer SG) becoming binding for high-income older members who cannot increase sacrifice without breaching the cap.",
                            "style": "amber",
                        },
                    },
                    {
                        "title": "Younger participants sacrifice a higher share of their salary",
                        "body": (
                            "When sacrifice is expressed as a percentage of annual salary, an inverse relationship with age emerges: "
                            "18–24 participants contribute 7.0% of salary, while 65+ participants contribute 5.1%. As salaries rise with age, "
                            "the $30K concessional cap (inclusive of 11.5% employer SG) constrains the effective sacrifice rate for older members "
                            "on higher incomes. A 55-year-old on $150,000 has only $12,750 of concessional headroom remaining after employer SG."
                        ),
                        "chart_path": "outputs/2026-04-29_sacrifice_pct_salary_by_age.png",
                        "chart_caption": "Figure 3: Salary sacrifice as % of salary — participants only. Declining rate reflects concessional cap effect for older, higher-paid cohorts.",
                    },
                ],
            },
            {
                "title": "Digital Engagement: What Does and Doesn't Predict Participation",
                "findings": [
                    {
                        "title": "App usage and login recency show no predictive signal",
                        "body": (
                            "Counter to expectations, digital engagement metrics are not positively associated with salary sacrifice participation. "
                            "In several age bands, members with the app installed participate at lower rates than those without. "
                            "Median days-since-last-login is nearly identical for participants and non-participants across all cohorts "
                            "(typically 29–39 days for both groups). The decision to salary sacrifice appears driven by financial capacity "
                            "and life-stage awareness — not platform engagement. This has implications for digital-first activation campaigns: "
                            "targeting based on app usage or login recency is unlikely to improve conversion."
                        ),
                        "table": {
                            "headers": ["Metric", "Participants (median/rate)", "Non-participants (median/rate)", "Signal?"],
                            "rows": [
                                ["Days since last login", "32 days", "34 days", "None"],
                                ["App installed rate", "43.2%", "47.8%", "None (negative)"],
                                ["Investment option: High Growth (18–24)", "16.2%", "7.1%", "Selection effect only"],
                            ],
                            "col_widths": [2.5, 1.8, 2.0, 1.0],
                            "caption": "Table 2: Digital engagement vs participation — no positive signal detected.",
                        },
                    },
                ],
            },
        ],

        recommendations=[
            {
                "title": "Launch age-triggered salary milestone nudges for 18–34 cohort",
                "body": (
                    "1,125 non-participating members aged 18–34. At lower salaries, participation is near-zero — "
                    "target members when they cross the ~$85K salary threshold with a proactive communication. "
                    "Focus on habit formation: a $50/fortnight start is better than no start. Consider a stepped-contribution "
                    "default (e.g. 1% of salary increase automatically directed to salary sacrifice)."
                ),
                "owner": "Member Experience",
                "timeframe": "Q3 2026",
                "priority": "medium",
            },
            {
                "title": "Deploy peer comparison nudge for 35–44 non-participants",
                "body": (
                    "718 non-participating members aged 35–44 represent the single largest activation opportunity. "
                    "48% of their cohort peers already participate — peer comparison messaging is well-evidenced. "
                    "Test a default opt-in mechanism (member must opt out rather than opt in) for this age band. "
                    "A/B test digital vs outbound call channel; this cohort may respond to both."
                ),
                "owner": "Member Experience + Digital",
                "timeframe": "Q2–Q3 2026",
                "priority": "high",
            },
            {
                "title": "Conduct outreach call program for 55–64 high-income non-participants",
                "body": (
                    "273 non-participating members in the 55–64 band include high-income earners ($129K median) "
                    "who may be at or near the concessional cap. Before assuming disengagement, investigate: "
                    "are they contributing via personal deductible contributions instead? Are they aware of "
                    "non-concessional options? A member service call will likely outperform digital for this cohort."
                ),
                "owner": "Member Services",
                "timeframe": "Q3 2026",
                "priority": "medium",
            },
            {
                "title": "Provide carry-forward concessional contribution guidance to 65+ participants",
                "body": (
                    "175 active participants aged 65+ are likely near the $30K concessional cap. "
                    "Members with TSB below $500K are eligible to carry forward unused concessional cap from prior years. "
                    "Proactive guidance on this mechanism — especially for members who took career breaks — "
                    "could unlock additional tax-effective contributions."
                ),
                "owner": "Member Advice + Communications",
                "timeframe": "Q4 2026",
                "priority": "low",
            },
        ],

        caveats=[
            "Cross-sectional snapshot (31 March 2026) only — causality cannot be established from observational data.",
            "Salary nulls: 100 members (2.0%) have missing salary data and are excluded from salary-stratified analyses.",
            "18–24 participant group is small (n=31) — conditional amounts for this band are indicative only.",
            "No adviser relationship data — likely an important omitted variable explaining residual participation variance.",
            "Concessional cap interactions not explicitly modelled — estimates of cap-binding effects are inferred, not measured.",
            "Regression model R²=9.5% — the model is explanatory, not predictive; significant unobserved variation remains.",
        ],

        appendix=[
            {
                "title": "Data quality flags",
                "type": "bullets",
                "content": [
                    "5,025 raw rows; 25 exact duplicates removed → 5,000 clean records",
                    "salary_aud: 100 nulls (2.0%) — excluded from salary analyses",
                    "snapshot_date: constant column (31/03/2026) — single point-in-time",
                    "salary_sacrifice_12mo_aud: median = $0 across all members (majority are non-participants)",
                    "No PII columns retained in outputs — member_id not included in any aggregation",
                ],
            },
            {
                "title": "Statistical methodology",
                "type": "bullets",
                "content": [
                    "Participation rate: binary flag (salary_sacrifice_12mo_aud > 0) / total members per band",
                    "Confidence intervals: Wilson score method for proportions (95% CI)",
                    "Conditional amounts: median and mean for participants only; IQR reported",
                    "Salary quartiles: computed within-fund (not national benchmarks)",
                    "OLS linear probability model: participation ~ age_rank + log(salary) + tenure_years",
                    "Standardised coefficients reported (mean-centred, unit-variance scaled)",
                    "Sacrifice as % of salary: salary_sacrifice / salary_aud, participants only",
                ],
            },
            {
                "title": "Generated artefacts",
                "type": "table",
                "content": {
                    "headers": ["File", "Type", "Description"],
                    "rows": [
                        ["_profile_super_members_2026Q1.md", "Markdown", "Automated data quality profile"],
                        ["2026-04-29_findings.json", "JSON", "Machine-readable results"],
                        ["2026-04-29_participation_rate_by_age_band.png", "Chart", "Participation rate by age band"],
                        ["2026-04-29_conditional_amount_by_age_band.png", "Chart", "Conditional sacrifice amounts"],
                        ["2026-04-29_participation_by_age_salary_quartile.png", "Chart", "Participation by age × salary quartile"],
                        ["2026-04-29_sacrifice_pct_salary_by_age.png", "Chart", "Sacrifice % of salary by age band"],
                        ["2026-04-29_salary_sacrifice_analysis.md", "Markdown", "Full written analysis report"],
                        ["2026-04-29_salary_sacrifice_report_v2.pptx", "PowerPoint", "Executive presentation (v2 design)"],
                        ["2026-04-29_salary_sacrifice_report_v2.docx", "Word", "This document (v2 design)"],
                    ],
                    "col_widths": [3.0, 1.0, 2.4],
                },
            },
        ],
    )

    out = Path("outputs/2026-04-29_salary_sacrifice_report_v2.docx")
    build_report(cfg, out)


if __name__ == "__main__":
    if len(sys.argv) > 1:
        build_from_json(sys.argv[1])
    else:
        build_salary_sacrifice_report()
